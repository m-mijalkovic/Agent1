from fastapi import FastAPI, UploadFile, File
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openai import AzureOpenAI
from dotenv import load_dotenv
from typing import List, Optional, TypedDict, Annotated
import os
from langchain_openai import AzureChatOpenAI, AzureOpenAIEmbeddings
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langgraph.prebuilt import create_react_agent
from langchain_core.tools import tool
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import DirectoryLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langgraph.graph import StateGraph, END
from typing_extensions import TypedDict
import operator
import docx
import io

# Load environment variables
load_dotenv()

# Initialize Azure OpenAI client
client = AzureOpenAI(
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION")
)

# Initialize LangChain Azure OpenAI client
langchain_llm = AzureChatOpenAI(
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    deployment_name=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
    temperature=0.7
)

# Define a demo tool for the agent
@tool
def get_weather(city: str) -> str:
    """Get the current weather for a specific city. Use this tool when the user asks about weather conditions."""
    # This is a demo tool with fake responses
    fake_weather_data = {
        "paris": "Sunny, 22°C with light clouds",
        "london": "Rainy, 15°C with heavy clouds",
        "new york": "Partly cloudy, 18°C with moderate wind",
        "tokyo": "Clear sky, 25°C with no wind",
        "sydney": "Sunny, 28°C with clear skies"
    }

    city_lower = city.lower()
    if city_lower in fake_weather_data:
        return f"Weather in {city}: {fake_weather_data[city_lower]}"
    else:
        return f"Weather in {city}: Sunny, 20°C (demo data - city not in database)"

# Create the agent with the tool
tools = [get_weather]

# Create the agent using langgraph's prebuilt create_react_agent
agent_executor = create_react_agent(langchain_llm, tools)

# Create a second LLM for validation
validator_llm = AzureChatOpenAI(
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    deployment_name=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
    temperature=0.3  # Lower temperature for more consistent validation
)

# Initialize embeddings for RAG
embeddings = AzureOpenAIEmbeddings(
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    azure_deployment=os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT", "text-embedding-3-small")
)

# Global vector store variable
vector_store = None

# RAG State for LangGraph workflow
class RAGState(TypedDict):
    question: str
    context: List[str]
    answer: str

# RAG workflow nodes
def retrieve_documents(state: RAGState) -> RAGState:
    """Retrieve relevant documents from vector store"""
    question = state["question"]
    if vector_store is None:
        state["context"] = ["No documents loaded in vector store."]
        return state

    # Retrieve top 3 most relevant documents
    docs = vector_store.similarity_search(question, k=3)
    state["context"] = [doc.page_content for doc in docs]
    return state

def generate_answer(state: RAGState) -> RAGState:
    """Generate answer using LLM with retrieved context"""
    question = state["question"]
    context = state["context"]

    # Create prompt with context
    context_str = "\n\n".join(context)
    prompt = f"""Based on the following context, answer the user's question. If the context doesn't contain enough information to answer the question, say so.

Context:
{context_str}

Question: {question}

Answer:"""

    response = langchain_llm.invoke([HumanMessage(content=prompt)])
    state["answer"] = response.content
    return state

# Create RAG workflow using LangGraph
workflow = StateGraph(RAGState)
workflow.add_node("retrieve", retrieve_documents)
workflow.add_node("generate", generate_answer)
workflow.add_edge("retrieve", "generate")
workflow.add_edge("generate", END)
workflow.set_entry_point("retrieve")
rag_chain = workflow.compile()

# Function to initialize vector store on startup
def initialize_vector_store():
    """Initialize an empty vector store for document uploads via UI"""
    global vector_store

    try:
        # Create an empty Chroma vector store
        # Using a dummy document to initialize, then we can add real documents via UI
        vector_store = Chroma(
            collection_name="company_docs",
            embedding_function=embeddings
        )

        print("Vector store initialized successfully (empty - ready for document uploads via UI)")

        # Optional: Load documents from folder if they exist (for backward compatibility)
        documents_path = "./documents"
        if os.path.exists(documents_path):
            try:
                loader = DirectoryLoader(documents_path, glob="**/*.txt", loader_cls=TextLoader)
                documents = loader.load()

                if documents:
                    # Split documents into chunks
                    text_splitter = RecursiveCharacterTextSplitter(
                        chunk_size=500,
                        chunk_overlap=50
                    )
                    splits = text_splitter.split_documents(documents)

                    # Add to vector store
                    vector_store.add_documents(splits)
                    print(f"Loaded {len(documents)} documents from folder, split into {len(splits)} chunks")
                else:
                    print("Documents folder exists but is empty - use UI to upload documents")
            except Exception as e:
                print(f"Could not load documents from folder (this is okay): {str(e)}")
                print("Use UI to upload documents")
        else:
            print("Documents folder not found - use UI to upload documents")

    except Exception as e:
        print(f"Error initializing vector store: {str(e)}")
        raise

app = FastAPI()

# Add CORS middleware to allow requests from the UI
# For Azure deployment, update allow_origins with your Azure Web App URL
# Example: ["https://your-app-name.azurewebsites.net"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, replace with specific origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")

# Serve the UI at the root
@app.get("/ui")
async def read_ui():
    return FileResponse("static/index.html")

# Startup event to initialize vector store
@app.on_event("startup")
async def startup_event():
    """Initialize vector store on application startup"""
    print("Initializing vector store...")
    initialize_vector_store()
    print("Vector store ready!")

# Message model for conversation history
class Message(BaseModel):
    role: str  # "user" or "assistant"
    content: str

# Request model for the OpenAI endpoint
class PromptRequest(BaseModel):
    prompt: str
    conversation_history: Optional[List[Message]] = None

@app.get("/")
async def read_root():
    """Serve the UI at the root path"""
    return FileResponse("static/index.html")

@app.get("/second")
def read_second():
    return {"message": "Second!!!"}

@app.post("/ask")
async def ask_openai(request: PromptRequest):
    try:
        # Build messages list from conversation history
        messages = []
        if request.conversation_history:
            messages = [{"role": msg.role, "content": msg.content} for msg in request.conversation_history]

        # Add the new user prompt
        messages.append({"role": "user", "content": request.prompt})

        # Call OpenAI with full conversation history
        response = client.chat.completions.create(
            model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            messages=messages
        )

        assistant_response = response.choices[0].message.content

        # Build updated conversation history
        updated_history = messages + [{"role": "assistant", "content": assistant_response}]

        return {
            "prompt": request.prompt,
            "response": assistant_response,
            "conversation_history": updated_history
        }
    except Exception as e:
        return {"error": str(e)}

@app.post("/ask-langchain")
async def ask_langchain(request: PromptRequest):
    try:
        # Build LangChain messages from conversation history
        langchain_messages = []
        if request.conversation_history:
            for msg in request.conversation_history:
                if msg.role == "user":
                    langchain_messages.append(HumanMessage(content=msg.content))
                elif msg.role == "assistant":
                    langchain_messages.append(AIMessage(content=msg.content))

        # Add the new user prompt
        langchain_messages.append(HumanMessage(content=request.prompt))

        # Call LangChain with full conversation history
        response = langchain_llm.invoke(langchain_messages)

        assistant_response = response.content

        # Build updated conversation history for the response
        updated_history = []
        if request.conversation_history:
            updated_history = [{"role": msg.role, "content": msg.content} for msg in request.conversation_history]

        updated_history.append({"role": "user", "content": request.prompt})
        updated_history.append({"role": "assistant", "content": assistant_response})

        return {
            "prompt": request.prompt,
            "response": assistant_response,
            "conversation_history": updated_history,
            "method": "langchain"
        }
    except Exception as e:
        return {"error": str(e)}

@app.post("/ask-agent")
async def ask_agent(request: PromptRequest):
    try:
        # Build message list for the agent
        messages = []
        if request.conversation_history:
            for msg in request.conversation_history:
                if msg.role == "user":
                    messages.append(HumanMessage(content=msg.content))
                elif msg.role == "assistant":
                    messages.append(AIMessage(content=msg.content))

        # Add the new user prompt
        messages.append(HumanMessage(content=request.prompt))

        # Invoke the agent with the messages
        result = agent_executor.invoke({"messages": messages})

        # Extract the final response
        assistant_response = result["messages"][-1].content

        # Build updated conversation history
        updated_history = []
        if request.conversation_history:
            updated_history = [{"role": msg.role, "content": msg.content} for msg in request.conversation_history]

        updated_history.append({"role": "user", "content": request.prompt})
        updated_history.append({"role": "assistant", "content": assistant_response})

        # Check if tools were used (messages will have tool calls if tools were used)
        tool_used = any(hasattr(msg, 'tool_calls') and msg.tool_calls for msg in result["messages"] if hasattr(msg, 'tool_calls'))

        return {
            "prompt": request.prompt,
            "response": assistant_response,
            "conversation_history": updated_history,
            "method": "agent",
            "tool_used": tool_used
        }
    except Exception as e:
        return {"error": str(e)}

@app.post("/ask-validated")
async def ask_validated(request: PromptRequest):
    try:
        max_retries = 3
        attempts = []

        for attempt in range(max_retries):
            # Step 1: Get response from first agent
            messages = []
            if request.conversation_history:
                for msg in request.conversation_history:
                    if msg.role == "user":
                        messages.append(HumanMessage(content=msg.content))
                    elif msg.role == "assistant":
                        messages.append(AIMessage(content=msg.content))

            messages.append(HumanMessage(content=request.prompt))

            # Invoke the first agent
            result = agent_executor.invoke({"messages": messages})
            first_agent_response = result["messages"][-1].content

            # Step 2: Validate the response with the second agent
            validation_prompt = f"""You are a validator. Your job is to check if the response correctly answers the user's request.

User Request: {request.prompt}

Agent Response: {first_agent_response}

Analyze if the response:
1. Directly addresses the user's question
2. Is accurate and complete
3. Is relevant to the request

Respond with ONLY one of these:
- "VALID" if the response is good
- "INVALID: [reason]" if the response needs improvement

Your validation:"""

            validation_result = validator_llm.invoke([HumanMessage(content=validation_prompt)])
            validation_response = validation_result.content.strip()

            # Store attempt information
            attempts.append({
                "attempt": attempt + 1,
                "response": first_agent_response,
                "validation": validation_response
            })

            # Step 3: Check validation result
            if validation_response.startswith("VALID"):
                # Build updated conversation history
                updated_history = []
                if request.conversation_history:
                    updated_history = [{"role": msg.role, "content": msg.content} for msg in request.conversation_history]

                updated_history.append({"role": "user", "content": request.prompt})
                updated_history.append({"role": "assistant", "content": first_agent_response})

                return {
                    "prompt": request.prompt,
                    "response": first_agent_response,
                    "conversation_history": updated_history,
                    "method": "validated-agent",
                    "validation_status": "PASSED",
                    "attempts": attempts,
                    "total_attempts": attempt + 1
                }

            # If not valid and not the last attempt, continue loop
            if attempt < max_retries - 1:
                # Add feedback to the prompt for next attempt
                request.prompt = f"{request.prompt}\n\nPrevious attempt was insufficient. Validator feedback: {validation_response}. Please provide a better response."

        # If we exhausted all retries, return the last response with failure status
        updated_history = []
        if request.conversation_history:
            updated_history = [{"role": msg.role, "content": msg.content} for msg in request.conversation_history]

        updated_history.append({"role": "user", "content": request.prompt})
        updated_history.append({"role": "assistant", "content": first_agent_response})

        return {
            "prompt": request.prompt,
            "response": first_agent_response,
            "conversation_history": updated_history,
            "method": "validated-agent",
            "validation_status": "FAILED_MAX_RETRIES",
            "attempts": attempts,
            "total_attempts": max_retries
        }

    except Exception as e:
        return {"error": str(e)}

@app.post("/ask-rag")
async def ask_rag(request: PromptRequest):
    """RAG endpoint that retrieves relevant documents and generates answer using LangGraph workflow"""
    try:
        if vector_store is None:
            return {"error": "Vector store not initialized. Please ensure documents are loaded."}

        # Run the RAG workflow
        result = rag_chain.invoke({
            "question": request.prompt,
            "context": [],
            "answer": ""
        })

        # Build conversation history
        updated_history = []
        if request.conversation_history:
            updated_history = [{"role": msg.role, "content": msg.content} for msg in request.conversation_history]

        updated_history.append({"role": "user", "content": request.prompt})
        updated_history.append({"role": "assistant", "content": result["answer"]})

        return {
            "prompt": request.prompt,
            "response": result["answer"],
            "conversation_history": updated_history,
            "method": "rag",
            "context_used": result["context"],
            "num_documents_retrieved": len(result["context"])
        }

    except Exception as e:
        return {"error": str(e)}

@app.post("/upload-document")
async def upload_document(file: UploadFile = File(...)):
    """Upload a txt or docx file and add it to the vector store"""
    try:
        # Check if file is a supported format
        if not (file.filename.endswith('.txt') or file.filename.endswith('.docx') or file.filename.endswith('.doc')):
            return {"error": "Only .txt, .doc, and .docx files are supported"}

        # Check if vector store exists
        if vector_store is None:
            return {"error": "Vector store not initialized. Please restart the application."}

        # Read file content
        content = await file.read()

        # Extract text based on file type
        if file.filename.endswith('.txt'):
            text_content = content.decode('utf-8')
        elif file.filename.endswith('.docx') or file.filename.endswith('.doc'):
            # Parse .docx file
            doc_file = docx.Document(io.BytesIO(content))
            # Extract text from all paragraphs
            paragraphs = [paragraph.text for paragraph in doc_file.paragraphs if paragraph.text.strip()]
            # Also extract text from tables
            for table in doc_file.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            text_content = '\n'.join(paragraphs)

            if not text_content.strip():
                return {"error": "No text content found in the document"}

        # Create a Document object
        doc = Document(
            page_content=text_content,
            metadata={"source": file.filename, "uploaded": True}
        )

        # Split document into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50
        )
        splits = text_splitter.split_documents([doc])

        # Add to existing vector store
        vector_store.add_documents(splits)

        return {
            "message": f"Successfully uploaded and processed {file.filename}",
            "filename": file.filename,
            "chunks_created": len(splits),
            "file_type": "text" if file.filename.endswith('.txt') else "word",
            "status": "success"
        }

    except UnicodeDecodeError:
        return {"error": "File encoding error. Please ensure the file is UTF-8 encoded."}
    except Exception as e:
        return {"error": f"Failed to process file: {str(e)}"}
